# app/services/prep_service.py

from docx import Document
from pathlib import Path
from datetime import datetime

# Chemin vers le template Word (placez votre visit_template.docx dans app/templates/)
TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / "visit_template.docx"

def build_docx(prep_id: str, client_data: dict, prompt: str) -> str:
    """
    Génère un rapport de visite Word à partir d'un template.
    - Remplace les placeholders {{CLIENT_NAME}}, {{KDNR}}, {{PREP_DATE}}.
    - Ajoute le prompt IA à la fin.
    - Sauvegarde sous /tmp/prep_docs/{prep_id}.docx.
    """
    # Prépare le répertoire de sortie
    output_dir = Path("/tmp/prep_docs")
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / f"{prep_id}.docx"

    # Charge le document template
    doc = Document(TEMPLATE_PATH)

    # Mapping des placeholders à remplacer
    now = datetime.utcnow().strftime("%Y-%m-%d %H:%M")
    mapping = {
        "{{CLIENT_NAME}}": client_data.get("name", ""),
        "{{KDNR}}": client_data.get("kdnr", ""),
        "{{PREP_DATE}}": now,
    }

    def replace_in_paragraphs():
        for para in doc.paragraphs:
            for placeholder, value in mapping.items():
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value))

    def replace_in_tables():
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in mapping.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

    # Applique les remplacements
    replace_in_paragraphs()
    replace_in_tables()

    # Ajout d'une nouvelle page pour le prompt IA
    doc.add_page_break()
    doc.add_heading("Prompt généré par l’IA", level=2)
    doc.add_paragraph(prompt)

    # Sauvegarde le document final
    doc.save(docx_path)

    return str(docx_path)
