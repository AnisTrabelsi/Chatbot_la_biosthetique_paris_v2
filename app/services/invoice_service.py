# app/services/invoice_service.py  (ou mieux: prep_service.py)
from docx import Document
from pathlib import Path

def build_docx(prep_id: str, client_data: dict, prompt: str) -> str:
    output_dir = Path("/tmp/prep_docs")
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / f"{prep_id}.docx"

    doc = Document()
    doc.add_heading(f"Rapport de visite: {client_data['name']}", level=1)
    doc.add_paragraph(f"Kdnr: {client_data['kdnr']}")
    doc.add_paragraph("Prompt généré par l’IA :")
    doc.add_paragraph(prompt)
    # Autres sections…
    doc.save(docx_path)

    return str(docx_path)
